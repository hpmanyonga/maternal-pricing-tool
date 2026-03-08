"""Generate pre-filled Medical Service Agreements (DOCX) for NOH Cash and Discovery."""

import hashlib
import math
from datetime import date, datetime, timedelta
from io import BytesIO
from pathlib import Path

from docx import Document

_DATA_DIR = Path(__file__).resolve().parent.parent / "data"
TEMPLATE_PATH = _DATA_DIR / "msa_template.docx"
DISCOVERY_TEMPLATE_PATH = _DATA_DIR / "msa_discovery_template.docx"


def _generate_reference(patient_name: str, id_number: str) -> str:
    """Short deterministic reference from patient details + date."""
    raw = f"{patient_name}|{id_number}|{datetime.utcnow().date()}"
    return "NOH-MSA-" + hashlib.sha256(raw.encode()).hexdigest()[:8].upper()


def _estimate_edd(gestational_age_weeks: float) -> date:
    """Estimate EDD from current GA (40 weeks total)."""
    remaining_weeks = 40.0 - gestational_age_weeks
    return date.today() + timedelta(weeks=remaining_weeks)


def generate_msa_docx(
    *,
    patient_name: str,
    id_number: str,
    gestational_age_weeks: float,
    global_fee: float,
    months_to_34_weeks: int,
    monthly_payment: float,
) -> BytesIO:
    """
    Fill in the MSA template and return an in-memory DOCX buffer.

    Parameters
    ----------
    patient_name : str
        Full name with title, e.g. "Mrs J Doe".
    id_number : str
        ID or passport number (and surname if different).
    gestational_age_weeks : float
        Current gestational age at booking.
    global_fee : float
        Total global maternity fee (from pricing engine).
    months_to_34_weeks : int
        Number of monthly instalments.
    monthly_payment : float
        Amount per month.

    Returns
    -------
    BytesIO
        In-memory DOCX ready for download.
    """
    doc = Document(str(TEMPLATE_PATH))

    # Table 0 is the letterhead — skip
    # Table 1 is Client Details
    details_table = doc.tables[1]

    field_map = {
        0: _generate_reference(patient_name, id_number),          # Reference Number
        1: patient_name,                                           # Title and Name
        2: id_number,                                              # ID/Passport Number and Surname
        3: date.today().strftime("%d %B %Y"),                      # Date of Agreement
        4: _estimate_edd(gestational_age_weeks).strftime("%d %B %Y"),  # EDD
        5: f"{gestational_age_weeks:.1f}",                         # Current Gestation
        6: f"R {global_fee:,.2f}",                                 # Global Maternity Fee
        7: str(months_to_34_weeks),                                # Settlement Time
        8: f"R {monthly_payment:,.2f}",                            # Monthly Installment
    }

    for row_idx, value in field_map.items():
        cell = details_table.rows[row_idx].cells[1]
        cell.text = value
        # Preserve the existing font by copying style from the label cell
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = details_table.rows[row_idx].cells[0].paragraphs[0].runs[0].font.size if details_table.rows[row_idx].cells[0].paragraphs[0].runs else None

    # Table 3 — Practice signature block: pre-fill name
    practice_table = doc.tables[3]
    practice_table.rows[0].cells[1].text = "Dr H P Manyonga"

    # Write to in-memory buffer
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# Plan-type display mapping for Discovery checkbox ticking
_PLAN_CHECKBOX_MAP = {
    "KEYCARE": "KeyCare",
    "SMART": "Smart",
    "COASTAL_ESSENTIAL": "Coastal/Essential",
    "CLASSIC": "Classic",
    "EXECUTIVE": "Executive",
}


def generate_discovery_msa_docx(
    *,
    title: str,
    first_name: str,
    surname: str,
    id_number: str,
    mobile: str,
    email: str,
    gestational_age_weeks: float,
    plan_type: str,
    membership_no: str,
    dependent_code: str,
    enrollment_route: str,
) -> BytesIO:
    """
    Fill in the Discovery Global Fee MSA template.

    No instalment fields — the patient is on the Discovery global fee,
    so the agreement simply records that they are part of the pilot.

    Returns an in-memory DOCX buffer.
    """
    doc = Document(str(DISCOVERY_TEMPLATE_PATH))

    # ------------------------------------------------------------------
    # Table 0 — Patient & Medical Aid Details (4-column: label, val, label, val)
    # ------------------------------------------------------------------
    t0 = doc.tables[0]
    ref = _generate_reference(f"{first_name} {surname}", id_number).replace("NOH-MSA-", "NOH-DIS-")
    edd = _estimate_edd(gestational_age_weeks)

    # Row 0: Reference Nr | Date of Agreement
    t0.rows[0].cells[1].text = ref
    t0.rows[0].cells[3].text = date.today().strftime("%d %B %Y")
    # Row 1: Title | Name
    t0.rows[1].cells[1].text = title
    t0.rows[1].cells[3].text = first_name
    # Row 2: ID/Passport | Surname
    t0.rows[2].cells[1].text = id_number
    t0.rows[2].cells[3].text = surname
    # Row 3: Mobile | Email
    t0.rows[3].cells[1].text = mobile
    t0.rows[3].cells[3].text = email
    # Row 4: EDD | Current Gestation
    t0.rows[4].cells[1].text = edd.strftime("%d %B %Y")
    t0.rows[4].cells[3].text = f"{gestational_age_weeks:.1f}"
    # Row 5: Medical Scheme | Plan
    t0.rows[5].cells[1].text = "Discovery Health"
    plan_display = _PLAN_CHECKBOX_MAP.get(plan_type, plan_type)
    t0.rows[5].cells[3].text = plan_display
    # Row 6: Membership No | Dependent Code
    t0.rows[6].cells[1].text = membership_no
    t0.rows[6].cells[3].text = dependent_code

    # ------------------------------------------------------------------
    # Table 1 — Programme Type: tick Discovery Global Fee + plan
    # ------------------------------------------------------------------
    t1 = doc.tables[1]
    # Tick Discovery row
    cell_disc = t1.rows[0].cells[0]
    for para in cell_disc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace("\u2610", "\u2611")  # ☐ → ☑

    # Tick the correct plan in the plan-type cell
    cell_plan = t1.rows[0].cells[1]
    for para in cell_plan.paragraphs:
        for run in para.runs:
            for pkey, plabel in _PLAN_CHECKBOX_MAP.items():
                if pkey == plan_type:
                    run.text = run.text.replace(f"\u2610 {plabel}", f"\u2611 {plabel}")

    # ------------------------------------------------------------------
    # Table 5 — Signatures: pre-fill practice name
    # ------------------------------------------------------------------
    t5 = doc.tables[5]
    t5.rows[2].cells[1].text = "Dr H P Manyonga"

    # ------------------------------------------------------------------
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
